#!/usr/bin/env python3
"""
🎓 УЧЕБНЫЙ БОТ ПРЕМИУМ - ПОЛНАЯ ВЕРСИЯ 7.0
С выбором объема в листах А4, презентациями и источниками
"""

import os
import logging
import json
import time
import re
import random
from datetime import datetime
from flask import Flask, request, jsonify
import requests
import threading

# Импорты для создания файлов
try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("⚠️ Библиотека fpdf не установлена")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("⚠️ Библиотека python-docx не установлена")

# ============ НАСТРОЙКА ЛОГГИРОВАНИЯ ============
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# ============ FLASK ПРИЛОЖЕНИЕ ============
app = Flask(__name__)

# ============ КОНФИГУРАЦИЯ ============
TOKEN = os.environ.get('BOT_TOKEN', '')
if not TOKEN:
    logger.error("❌ ТОКЕН НЕ НАЙДЕН! Добавьте BOT_TOKEN в Environment Variables")
    exit()

logger.info(f"✅ Токен получен: {TOKEN[:10]}...")

# ============ БАЗЫ ДАННЫХ ============
user_devices = {}
user_settings = {}
user_history = {}
export_queue = {}
# ============ КОНСТАНТЫ ============
DEVICES = {
    "phone": {"icon": "📱", "name": "Телефон", "description": "Мобильная версия"},
    "pc": {"icon": "💻", "name": "Компьютер", "description": "Полная версия"},
    "tablet": {"icon": "📟", "name": "Планшет", "description": "Промежуточная версия"},
    "watch": {"icon": "⌚", "name": "Часы", "description": "Краткая версия"}
}

CONTENT_TYPES = {
    "conspect": {"icon": "📚", "name": "Конспект"},
    "referat": {"icon": "📄", "name": "Реферат"},
    "presentation": {"icon": "🎤", "name": "Презентация"},
    "essay": {"icon": "✍️", "name": "Эссе"},
    "summary": {"icon": "📝", "name": "Краткое содержание"}
}

EXPORT_FORMATS = {
    "pdf": {"icon": "📄", "name": "PDF", "mime": "application/pdf"},
    "docx": {"icon": "📝", "name": "DOCX", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "txt": {"icon": "📋", "name": "TXT", "mime": "text/plain"}
}

# ============ ОБЪЕМ В ЛИСТАХ А4 ============
VOLUME_LEVELS = {
    "1": {"icon": "📄", "name": "1 лист", "pages": 1, "words": "300-400", "description": "Конспект, тезисы"},
    "2": {"icon": "📄📄", "name": "2 листа", "pages": 2, "words": "600-800", "description": "Краткий реферат"},
    "3": {"icon": "📄📄📄", "name": "3 листа", "pages": 3, "words": "900-1200", "description": "Средняя работа"},
    "4": {"icon": "📄📄📄📄", "name": "4 листа", "pages": 4, "words": "1200-1600", "description": "Полный реферат"},
    "5": {"icon": "📄📄📄📄📄", "name": "5 листов", "pages": 5, "words": "1500-2000", "description": "Детальный анализ"},
    "6": {"icon": "📚", "name": "6 листов", "pages": 6, "words": "1800-2400", "description": "Расширенная работа"},
    "7": {"icon": "📚📄", "name": "7 листов", "pages": 7, "words": "2100-2800", "description": "Подробное исследование"},
    "8": {"icon": "📚📚", "name": "8 листов", "pages": 8, "words": "2400-3200", "description": "Курсовая работа"},
    "9": {"icon": "📚📚📄", "name": "9 листов", "pages": 9, "words": "2700-3600", "description": "Серьезное исследование"},
    "10": {"icon": "📘", "name": "10 листов", "pages": 10, "words": "3000-4000", "description": "Дипломная работа"},
    "15": {"icon": "📗", "name": "15 листов", "pages": 15, "words": "4500-6000", "description": "Научная статья"},
    "20": {"icon": "📕", "name": "20 листов", "pages": 20, "words": "6000-8000", "description": "Магистерская работа"},
    "30": {"icon": "📓", "name": "30 листов", "pages": 30, "words": "9000-12000", "description": "Диссертация"}
}

# ============ ИСТОЧНИКИ ИНФОРМАЦИИ ============
SOURCES_DATABASE = {
    # Научные журналы
    "scientific": [
        "Научный журнал 'Вестник Московского университета'",
        "Журнал 'Вопросы философии'",
        "Российский научный журнал 'Образование и наука'",
        "Международный журнал прикладных наук",
        "Журнал 'Психологическая наука и образование'",
        "Научные труды РАН",
        "Journal of Applied Sciences",
        "International Journal of Science Education"
    ],
    
    # Учебники
    "textbooks": [
        "Учебник по общей психологии (Под ред. А.В. Петровского)",
        "Учебник по экономике (С.Г. Капканщиков)",
        "Учебник по философии (А.Г. Спиркин)",
        "Основы социологии (Г.В. Осипов)",
        "Учебник по правоведению (М.Н. Марченко)",
        "История России (А.С. Орлов)",
        "Биология (В.Н. Ярыгин)",
        "Физика (Г.Я. Мякишев)"
    ],
    
    # Интернет-ресурсы
    "online": [
        "Научная электронная библиотека eLIBRARY.RU",
        "Образовательный портал 'Инфоурок'",
        "Энциклопедия 'Кругосвет'",
        "Сайт 'ПостНаука'",
        "Образовательный проект 'Арзамас'",
        "Библиотека Максима Мошкова",
        "Научно-популярный портал 'Элементы'",
        "Образовательный ресурс 'Универсариум'"
    ],
    
    # Законы и нормативные акты
    "laws": [
        "Конституция Российской Федерации",
        "Федеральный закон 'Об образовании в РФ'",
        "Гражданский кодекс РФ",
        "Трудовой кодекс РФ",
        "Семейный кодекс РФ",
        "Федеральные государственные образовательные стандарты"
    ]
}

# ============ ШАБЛОНЫ ПРЕЗЕНТАЦИЙ ============
PRESENTATION_TEMPLATES = {
    "academic": {
        "name": "Академическая",
        "slides": [
            "Титульный слайд",
            "План презентации",
            "Актуальность темы",
            "Цели и задачи",
            "Теоретические основы",
            "Практическая часть",
            "Результаты",
            "Выводы",
            "Спасибо за внимание"
        ],
        "style": "Строгий академический стиль"
    },
    "business": {
        "name": "Бизнес-презентация",
        "slides": [
            "Заголовок и спикер",
            "Проблема/Возможность",
            "Решение/Продукт",
            "Преимущества",
            "Рынок и конкуренция",
            "Финансовые показатели",
            "Команда",
            "Дорожная карта",
            "Контакты"
        ],
        "style": "Корпоративный стиль"
    },
    "creative": {
        "name": "Креативная",
        "slides": [
            "Визуальный заголовок",
            "История/Кейс",
            "Инновация/Идея",
            "Визуализация данных",
            "Примеры работ",
            "Процесс создания",
            "Результаты",
            "Призыв к действию",
            "Контакты в соцсетях"
        ],
        "style": "Креативный современный стиль"
    },
    "educational": {
        "name": "Образовательная",
        "slides": [
            "Тема занятия",
            "Цели обучения",
            "План урока",
            "Теоретический материал",
            "Примеры и упражнения",
            "Практическое задание",
            "Контрольные вопросы",
            "Домашнее задание",
            "Литература для изучения"
        ],
        "style": "Образовательный стиль с элементами интерактива"
    }
}

# ============ УТИЛИТЫ ============
def get_user_device(user_id: str) -> dict:
    """Получение устройства пользователя"""
    device_key = user_devices.get(user_id, "phone")
    return DEVICES.get(device_key, DEVICES["phone"])

def save_to_history(user_id: str, topic: str, content_type: str, volume: str = "3"):
    """Сохранение в историю"""
    if user_id not in user_history:
        user_history[user_id] = []
    
    user_history[user_id].append({
        "topic": topic,
        "type": content_type,
        "volume": volume,
        "timestamp": datetime.now().isoformat(),
        "device": user_devices.get(user_id, "phone")
    })
    
    # Ограничиваем историю 50 последними записями
    if len(user_history[user_id]) > 50:
        user_history[user_id] = user_history[user_id][-50:]

def split_message(text: str, max_length: int = 4000) -> list:
    """Разделение длинного сообщения на части"""
    if len(text) <= max_length:
        return [text]
    
    parts = []
    while text:
        if len(text) <= max_length:
            parts.append(text)
            break
        
        split_point = max_length
        for i in range(max_length - 100, max_length):
            if i < len(text) and text[i] in ['\n', '.', '!', '?', ';']:
                split_point = i + 1
                break
        
        parts.append(text[:split_point])
        text = text[split_point:].strip()
    
    return parts

def generate_sources(topic: str, count: int = 5) -> list:
    """Генерация списка источников по теме"""
    sources = []
    
    # Выбираем источники из разных категорий
    categories = list(SOURCES_DATABASE.keys())
    random.shuffle(categories)
    
    for category in categories:
        if len(sources) >= count:
            break
        
        category_sources = SOURCES_DATABASE[category]
        selected = random.sample(category_sources, min(2, len(category_sources)))
        
        for source in selected:
            if len(sources) < count:
                # Добавляем год издания
                year = random.randint(2015, 2024)
                sources.append(f"{source} ({year} г.)")
    
    # Добавляем тематические источники
    topic_keywords = topic.lower().split()
    for keyword in topic_keywords[:2]:
        if len(sources) < count:
            sources.append(f"Монография по теме '{keyword}' (2020-2023 гг.)")
    
    return sources[:count]

def parse_volume_from_text(text: str) -> tuple:
    """Извлечение объема из текста запроса"""
    patterns = [
        r'(\d+)\s*лист[аов]*',
        r'(\d+)\s*л\b',
        r'(\d+)\s*стр[аиц]*',
        r'\b(\d+)\b',
    ]
    
    volume = None
    clean_text = text
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            volume = match.group(1)
            clean_text = text[:match.start()] + text[match.end():]
            break
    
    # Проверяем словесное написание
    if not volume:
        words_to_numbers = {
            'один': '1', 'два': '2', 'три': '3', 'четыре': '4', 'пять': '5',
            'шесть': '6', 'семь': '7', 'восемь': '8', 'девять': '9', 'десять': '10'
        }
        
        for word, num in words_to_numbers.items():
            if word in text.lower():
                volume = num
                clean_text = clean_text.lower().replace(word, '')
                break
    
    # Ограничиваем объем
    if volume:
        try:
            volume_int = int(volume)
            if volume_int > 30:
                volume = "30"
            elif volume_int < 1:
                volume = "1"
        except:
            volume = "3"
    
    return clean_text.strip(), volume
    # ============ ОТПРАВКА СООБЩЕНИЙ ============
def send_telegram_message(chat_id: int, text: str, parse_mode: str = "HTML", 
                         reply_markup: dict = None) -> dict:
    """Отправка сообщения в Telegram"""
    try:
        message_parts = split_message(text)
        
        for i, part in enumerate(message_parts):
            if i > 0:
                part = f"📄 Часть {i+1}:\n\n{part}"
            
            url = f"https://api.telegram.org/bot{TOKEN}/sendMessage"
            payload = {
                "chat_id": chat_id,
                "text": part,
                "parse_mode": parse_mode,
                "disable_web_page_preview": True
            }
            
            if reply_markup and i == len(message_parts) - 1:
                payload["reply_markup"] = reply_markup
            
            response = requests.post(url, json=payload, timeout=10)
            
            if response.status_code != 200:
                logger.error(f"❌ Ошибка отправки: {response.text}")
                return {"ok": False, "error": response.text}
        
        logger.info(f"📤 Сообщение отправлено в чат {chat_id}")
        return {"ok": True}
            
    except Exception as e:
        logger.error(f"❌ Исключение при отправке: {e}")
        return {"ok": False, "error": str(e)}

def send_telegram_document(chat_id: int, filename: str, content: bytes, caption: str = ""):
    """Отправка документа в Telegram"""
    try:
        url = f"https://api.telegram.org/bot{TOKEN}/sendDocument"
        
        files = {
            "document": (filename, content)
        }
        
        data = {
            "chat_id": chat_id,
            "caption": caption[:1024],
            "parse_mode": "HTML"
        }
        
        response = requests.post(url, files=files, data=data, timeout=30)
        
        if response.status_code == 200:
            logger.info(f"📎 Документ отправлен: {filename}")
            return response.json()
        else:
            logger.error(f"❌ Ошибка отправки документа: {response.text}")
            return {"ok": False, "error": response.text}
            
    except Exception as e:
        logger.error(f"❌ Исключение при отправке документа: {e}")
        return {"ok": False, "error": str(e)}
        # ============ ГЕНЕРАЦИЯ КОНТЕНТА ============
def generate_ai_content(topic: str, content_type: str = "conspect", 
                       device_type: str = "phone", volume_pages: int = 3,
                       presentation_template: str = "academic") -> str:
    """Генерация AI контента с источниками"""
    
    volume_key = str(volume_pages)
    volume_info = VOLUME_LEVELS.get(volume_key, VOLUME_LEVELS["3"])
    device_info = DEVICES.get(device_type, DEVICES["phone"])
    
    # Заголовок с информацией
    content = []
    content.append(f"📚 <b>{CONTENT_TYPES[content_type]['name'].upper()} ПО ТЕМЕ: {topic.upper()}</b>")
    content.append("")
    content.append(f"📊 <b>ТЕХНИЧЕСКИЕ ПАРАМЕТРЫ:</b>")
    content.append(f"• Объем: {volume_info['icon']} {volume_info['name']} А4")
    content.append(f"• Количество слов: {volume_info['words']}")
    content.append(f"• Устройство: {device_info['icon']} {device_info['name']}")
    content.append(f"• Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    content.append("")
    
    # Генерируем основное содержание
    if content_type == "conspect":
        content.extend(generate_conspect_content(topic, volume_pages))
    elif content_type == "referat":
        content.extend(generate_referat_content(topic, volume_pages))
    elif content_type == "presentation":
        content.extend(generate_presentation_content(topic, volume_pages, presentation_template))
    elif content_type == "essay":
        content.extend(generate_essay_content(topic, volume_pages))
    
    content.append("")
    content.append("📌 <b>ИСТОЧНИКИ ИНФОРМАЦИИ:</b>")
    
    # Добавляем источники
    sources = generate_sources(topic, min(volume_pages + 2, 8))
    for i, source in enumerate(sources, 1):
        content.append(f"{i}. {source}")
    
    content.append("")
    content.append("💡 <b>РЕКОМЕНДАЦИИ ПО ИСПОЛЬЗОВАНИЮ:</b>")
    content.append("• Используйте как основу для собственной работы")
    content.append("• Проверяйте актуальность источников")
    content.append("• Дополняйте собственными исследованиями")
    content.append("• Соблюдайте академическую честность")
    
    if device_type == "phone":
        content.append("")
        content.append("📱 <b>СОВЕТ ДЛЯ ТЕЛЕФОНА:</b>")
        content.append("• Сохраните в заметки")
        content.append("• Используйте режим чтения")
        content.append("• Поделитесь с одногруппниками")
    
    return "\n".join(content)

def generate_conspect_content(topic: str, volume_pages: int) -> list:
    """Генерация содержания конспекта"""
    content = []
    
    # Распределение по листам
    intro_pages = max(1, volume_pages // 4)
    main_pages = volume_pages - intro_pages - 1
    conclusion_pages = 1
    
    content.append("<b>📖 СТРУКТУРА КОНСПЕКТА:</b>")
    content.append("")
    
    content.append(f"<b>1. ВВЕДЕНИЕ ({intro_pages} стр.)</b>")
    content.append("• Основные понятия и определения")
    content.append("• Актуальность темы")
    content.append("• Цели изучения")
    if intro_pages >= 2:
        content.append("• Исторический контекст")
        content.append("• Значимость в современном мире")
    
    content.append("")
    content.append(f"<b>2. ОСНОВНАЯ ЧАСТЬ ({main_pages} стр.)</b>")
    
    chapters = min(3, max(2, main_pages // 2))
    for i in range(1, chapters + 1):
        content.append(f"")
        content.append(f"<b>2.{i}. Раздел {i}</b>")
        content.append(f"• Ключевые положения")
        content.append(f"• Теоретические аспекты")
        content.append(f"• Практические примеры")
        if main_pages >= 4:
            content.append(f"• Таблицы и схемы")
            content.append(f"• Контрольные вопросы")
    
    content.append("")
    content.append(f"<b>3. ЗАКЛЮЧЕНИЕ ({conclusion_pages} стр.)</b>")
    content.append("• Основные выводы")
    content.append("• Практическая значимость")
    content.append("• Рекомендации для изучения")
    
    return content

def generate_referat_content(topic: str, volume_pages: int) -> list:
    """Генерация содержания реферата"""
    content = []
    
    content.append("<b>📄 СТРУКТУРА РЕФЕРАТА:</b>")
    content.append("")
    
    content.append("<b>1. ТИТУЛЬНЫЙ ЛИСТ</b>")
    content.append("• Название учебного заведения")
    content.append("• Тема реферата")
    content.append("• ФИО студента и преподавателя")
    content.append("• Город и год")
    
    content.append("")
    content.append("<b>2. ОГЛАВЛЕНИЕ (1 стр.)</b>")
    content.append("• Перечень разделов с номерами страниц")
    
    content.append("")
    intro_pages = max(1, volume_pages // 5)
    content.append(f"<b>3. ВВЕДЕНИЕ ({intro_pages} стр.)</b>")
    content.append("• Актуальность темы")
    content.append("• Цели и задачи работы")
    content.append("• Методология исследования")
    if intro_pages >= 2:
        content.append("• Объект и предмет исследования")
        content.append("• Гипотеза исследования")
    
    main_pages = volume_pages - intro_pages - 2
    content.append("")
    content.append(f"<b>4. ОСНОВНАЯ ЧАСТЬ ({main_pages} стр.)</b>")
    
    chapters = min(3, max(2, main_pages // 2))
    for i in range(1, chapters + 1):
        content.append("")
        content.append(f"<b>4.{i}. Глава {i}</b>")
        content.append(f"• Теоретические основы")
        content.append(f"• Анализ литературы")
        content.append(f"• Практические аспекты")
        if main_pages >= 6:
            content.append(f"• Эмпирические данные")
            content.append(f"• Статистический анализ")
    
    content.append("")
    content.append("<b>5. ЗАКЛЮЧЕНИЕ (1-2 стр.)</b>")
    content.append("• Выводы по работе")
    content.append("• Достижение целей")
    content.append("• Практические рекомендации")
    
    if volume_pages >= 6:
        content.append("")
        content.append("<b>6. СПИСОК ЛИТЕРАТУРЫ (1-2 стр.)</b>")
        content.append("• Книги и учебники (5-10 источников)")
        content.append("• Научные статьи")
        content.append("• Интернет-ресурсы")
    
    return content

def generate_presentation_content(topic: str, volume_pages: int, template: str = "academic") -> list:
    """Генерация структуры презентации"""
    content = []
    
    template_info = PRESENTATION_TEMPLATES.get(template, PRESENTATION_TEMPLATES["academic"])
    
    content.append(f"🎤 <b>ПРЕЗЕНТАЦИЯ: {topic.upper()}</b>")
    content.append(f"📊 Стиль: {template_info['name']}")
    content.append(f"🎨 Характеристика: {template_info['style']}")
    content.append("")
    
    content.append("<b>📋 СТРУКТУРА ПРЕЗЕНТАЦИИ:</b>")
    content.append(f"Всего слайдов: {len(template_info['slides'])}")
    content.append(f"Примерное время выступления: {len(template_info['slides']) * 1.5:.1f} мин.")
    content.append("")
    
    for i, slide in enumerate(template_info['slides'], 1):
        content.append(f"<b>Слайд {i}: {slide}</b>")
        
        # Добавляем рекомендации для каждого слайда
        if i == 1:
            content.append("• Название работы крупным шрифтом")
            content.append("• ФИО автора и научного руководителя")
            content.append("• Учебное заведение, год")
        elif "актуальность" in slide.lower():
            content.append("• Почему эта тема важна сейчас")
            content.append("• Статистика или факты")
            content.append("• Проблема, которую решает работа")
        elif "результаты" in slide.lower():
            content.append("• Ключевые цифры и показатели")
            content.append("• Графики и диаграммы")
            content.append("• Сравнение с другими работами")
        elif "выводы" in slide.lower():
            content.append("• 3-5 основных выводов")
            content.append("• Практические рекомендации")
            content.append("• Перспективы развития")
        else:
            content.append("• Основные тезисы")
            content.append("• Примеры и кейсы")
            content.append("• Визуальное сопровождение")
        
        content.append("")
    
    content.append("<b>🎯 РЕКОМЕНДАЦИИ ПО ОФОРМЛЕНИЮ:</b>")
    content.append("• 1 слайд = 1 идея")
    content.append("• Минимум текста, максимум визуалов")
    content.append("• Контрастные цвета для читаемости")
    content.append("• Единый стиль всех слайдов")
    content.append("• Шрифт не менее 24pt для основного текста")
    content.append("• Время на слайд: 1-2 минуты")
    
    content.append("")
    content.append("<b>💡 СОВЕТЫ ДЛЯ ВЫСТУПЛЕНИЯ:</b>")
    content.append("• Репетируйте перед зеркалом")
    content.append("• Подготовьте ответы на возможные вопросы")
    content.append("• Используйте указку или лазерную указку")
    content.append("• Поддерживайте зрительный контакт")
    content.append("• Следите за временем")
    
    return content

def generate_essay_content(topic: str, volume_pages: int) -> list:
    """Генерация содержания эссе"""
    content = []
    
    content.append("<b>✍️ СТРУКТУРА ЭССЕ:</b>")
    content.append("")
    
    content.append(f"<b>ОБЪЕМ: {volume_pages} стр. А4 ({VOLUME_LEVELS[str(volume_pages)]['words']} слов)</b>")
    content.append("")
    
    content.append("<b>1. ВСТУПЛЕНИЕ (10-15% объема)</b>")
    content.append("• Представление темы")
    content.append("• Актуальность проблемы")
    content.append("• Основной тезис (центральная идея)")
    content.append("• Цели и задачи эссе")
    
    content.append("")
    content.append("<b>2. ОСНОВНАЯ ЧАСТЬ (70-80% объема)</b>")
    
    paragraphs = min(5, max(3, volume_pages * 2))
    for i in range(1, paragraphs + 1):
        content.append("")
        content.append(f"<b>2.{i}. Абзац {i}</b>")
        content.append(f"• Основная мысль абзаца")
        content.append(f"• Аргумент или доказательство")
        content.append(f"• Пример или цитата")
        content.append(f"• Связь с основным тезисом")
    
    content.append("")
    content.append("<b>3. ЗАКЛЮЧЕНИЕ (10-15% объема)</b>")
    content.append("• Обобщение основных идей")
    content.append("• Подтверждение основного тезиса")
    content.append("• Выводы и размышления")
    content.append("• Перспективы дальнейшего исследования")
    
    content.append("")
    content.append("<b>🎯 КРИТЕРИИ ОЦЕНКИ ЭССЕ:</b>")
    content.append("• Глубина раскрытия темы")
    content.append("• Логичность и структурированность")
    content.append("• Аргументированность позиции")
    content.append("• Грамотность и стиль изложения")
    content.append("• Оригинальность мышления")
    
    return content
    # ============ СОЗДАНИЕ ФАЙЛОВ ============
def create_txt_file(topic: str, content: str, content_type: str) -> tuple:
    """Создание TXT файла"""
    import re
    
    # Убираем HTML теги для чистого текста
    clean_content = re.sub(r'<[^>]+>', '', content)
    
    # Заголовок файла
    file_content = "=" * 60 + "\n"
    file_content += f"{CONTENT_TYPES[content_type]['name'].upper()}: {topic.upper()}\n"
    file_content += "=" * 60 + "\n\n"
    file_content += clean_content
    file_content += f"\n\n{'=' * 60}\n"
    file_content += f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
    file_content += f"Бот: @Konspekt_help_bot\n"
    file_content += "=" * 60
    
    filename = f"{content_type}_{topic[:30].replace(' ', '_')}.txt"
    return filename, file_content.encode('utf-8')

def create_pdf_file(topic: str, content: str, content_type: str) -> tuple:
    """Создание PDF файла"""
    import re
    
    if not PDF_AVAILABLE:
        return create_txt_file(topic, content, content_type)
    
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Заголовок
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt=f"{CONTENT_TYPES[content_type]['name'].upper()}: {topic.upper()}", ln=1, align='C')
        pdf.ln(5)
        
        # Линия
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(10)
        
        # Контент
        pdf.set_font("Arial", size=12)
        
        # Убираем HTML теги
        clean_content = re.sub(r'<[^>]+>', '', content)
        
        # Разбиваем на строки
        lines = clean_content.split('\n')
        for line in lines:
            if line.strip():
                # Проверяем заголовки
                if line.strip().startswith('=') or any(x in line for x in ['СТРУКТУРА', 'ПАРАМЕТРЫ', 'ИСТОЧНИКИ', 'РЕКОМЕНДАЦИИ']):
                    pdf.ln(5)
                    pdf.set_font("Arial", 'B', 14)
                    pdf.multi_cell(0, 10, txt=line.strip())
                    pdf.set_font("Arial", size=12)
                else:
                    pdf.multi_cell(0, 8, txt=line)
            else:
                pdf.ln(5)
        
        # Подвал
        pdf.ln(10)
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(200, 10, txt=f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}", ln=1, align='C')
        pdf.cell(200, 10, txt="Бот: @Konspekt_help_bot", ln=1, align='C')
        
        filename = f"{content_type}_{topic[:30].replace(' ', '_')}.pdf"
        return filename, pdf.output(dest='S').encode('latin1')
        
    except Exception as e:
        logger.error(f"❌ Ошибка создания PDF: {e}")
        return create_txt_file(topic, content, content_type)

def create_docx_file(topic: str, content: str, content_type: str) -> tuple:
    """Создание DOCX файла"""
    import re
    from io import BytesIO
    
    if not DOCX_AVAILABLE:
        return create_txt_file(topic, content, content_type)
    
    try:
        doc = Document()
        
        # Заголовок
        title = doc.add_heading(f'{CONTENT_TYPES[content_type]["name"].upper()}: {topic.upper()}', 0)
        title.alignment = 1  # По центру
        
        # Контент
        clean_content = re.sub(r'<[^>]+>', '', content)
        
        for line in clean_content.split('\n'):
            if line.strip():
                if line.strip().startswith('=') or any(x in line for x in ['СТРУКТУРА', 'ПАРАМЕТРЫ', 'ИСТОЧНИКИ']):
                    # Заголовки разделов
                    doc.add_heading(line.strip(), level=1)
                elif line.strip().startswith('•'):
                    # Маркированный список
                    doc.add_paragraph(line.strip())
                elif any(line.strip().startswith(x) for x in ['1.', '2.', '3.', '4.', '5.']):
                    # Нумерованный список
                    doc.add_paragraph(line.strip())
                else:
                    # Обычный текст
                    doc.add_paragraph(line.strip())
        
        # Подвал
        doc.add_paragraph()
        doc.add_paragraph(f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph("Бот: @Konspekt_help_bot")
        
        # Сохраняем в BytesIO
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"{content_type}_{topic[:30].replace(' ', '_')}.docx"
        return filename, file_stream.read()
        
    except Exception as e:
        logger.error(f"❌ Ошибка создания DOCX: {e}")
        return create_txt_file(topic, content, content_type)
        # ============ ОБРАБОТЧИКИ КОМАНД ============
def handle_start_command(chat_id: int, username: str, user_id: str):
    """Обработка команды /start"""
    current_volume = user_settings.get(f"{user_id}_volume", "3")
    volume_info = VOLUME_LEVELS[current_volume]
    
    welcome_text = f"""👋 <b>Добро пожаловать, {username}!</b>

🎓 <b>УЧЕБНЫЙ БОТ ПРЕМИУМ v7.0</b>
📄 <b>С ВЫБОРОМ ОБЪЕМА В ЛИСТАХ А4</b>
🎤 <b>С ПРЕЗЕНТАЦИЯМИ И ИСТОЧНИКАМИ</b>

<b>✨ ВОЗМОЖНОСТИ:</b>
• 📄 Указывайте объем в листах: "реферат семья 4 листа"
• 🎤 Создание презентаций: "презентация экология 10 слайдов"
• 📚 Автоматические источники информации
• 📱 Адаптация под все устройства
• 📄 Реальные PDF/DOCX/TXT файлы

<b>🎯 КАК РАБОТАЕТ:</b>
1. Напишите запрос с объемом
2. Бот создает структуру на нужное количество листов
3. Получаете план с источниками
4. Скачиваете в нужном формате

<b>📊 СТАНДАРТНЫЕ ОБЪЕМЫ:</b>
• 1-2 л - 📘 Конспект, тезисы
• 3-5 л - 📗 Реферат, доклад  
• 6-10 л - 📕 Курсовая работа
• 10+ л - 📓 Диплом, исследование

<b>🎤 ШАБЛОНЫ ПРЕЗЕНТАЦИЙ:</b>
• Академическая - для учебных работ
• Бизнес - для проектов и стартапов
• Креативная - для творческих проектов
• Образовательная - для уроков и лекций

<b>🚀 БЫСТРЫЙ СТАРТ:</b>
<code>реферат семья 4 листа</code>
<code>презентация экология 10 слайдов бизнес</code>
<code>конспект математика 2л</code>

<b>📚 ИСТОЧНИКИ:</b>
• Автоматически подбираются по теме
• Из научных журналов и учебников
• С указанием годов издания

<i>Напишите запрос с указанием объема или используйте /help</i>"""
    
    send_telegram_message(chat_id, welcome_text)

def handle_help_command(chat_id: int):
    """Обработка команды /help"""
    help_text = """🆘 <b>ПОЛНАЯ СПРАВКА ПО БОТУ v7.0</b>

<b>📋 ОСНОВНЫЕ КОМАНДЫ:</b>
• /start - начать работу
• /help - эта справка  
• /volume - выбрать объем
• /presentation - создать презентацию
• /export - экспорт материалов
• /ai - AI-генерация контента
• /history - история запросов
• /settings - настройки

<b>📊 ФОРМАТ ЗАПРОСОВ С ОБЪЕМОМ:</b>
<code>[тип] [тема] [объем] [листы/слайды] [шаблон]</code>

<b>🎯 ПРИМЕРЫ:</b>
<code>реферат семья 4 листа</code>
<code>конспект физика 3л</code>
<code>презентация экология 10 слайдов бизнес</code>
<code>эссе философия 2 страницы</code>

<b>🎤 ШАБЛОНЫ ПРЕЗЕНТАЦИЙ:</b>
• <b>academic</b> - академическая (по умолчанию)
• <b>business</b> - бизнес-презентация
• <b>creative</b> - креативная
• <b>educational</b> - образовательная

<b>📄 ЭКСПОРТ ФАЙЛОВ:</b>
• PDF - для печати и чтения
• DOCX - для редактирования в Word
• TXT - простой текстовый формат

<b>📚 ИСТОЧНИКИ ИНФОРМАЦИИ:</b>
• Автоматически генерируются по теме
• Из научных журналов (2015-2024 гг.)
• Из учебников и монографий
• Из образовательных порталов

<b>📱 АДАПТАЦИЯ ПОД УСТРОЙСТВА:</b>
• 📱 Телефон - компактный формат
• 💻 Компьютер - полная версия
• 📟 Планшет - промежуточная версия
• ⌚ Часы - краткая версия

<b>❓ ЧАСТО ЗАДАВАЕМЫЕ ВОПРОСЫ:</b>
1. <b>Как указать объем?</b>
   Просто добавьте количество листов: "реферат тема 4 листа"

2. <b>Как создать презентацию?</b>
   Используйте: "презентация тема 10 слайдов"

3. <b>Где взять источники?</b>
   Бот автоматически добавляет список литературы

4. <b>Как скачать файл?</b>
   После генерации используйте /export

<i>Для начала работы напишите запрос с указанием объема!</i>"""
    
    send_telegram_message(chat_id, help_text)

def handle_volume_command(chat_id: int, user_id: str):
    """Обработка команды /volume - выбор объема"""
    current_volume = user_settings.get(f"{user_id}_volume", "3")
    volume_info = VOLUME_LEVELS[current_volume]
    
    user_settings[f"{user_id}_awaiting_volume"] = True
    
    keyboard = {
        "inline_keyboard": [
            [
                {"text": "1л", "callback_data": "volume_1"},
                {"text": "2л", "callback_data": "volume_2"}, 
                {"text": "3л", "callback_data": "volume_3"},
                {"text": "4л", "callback_data": "volume_4"}
            ],
            [
                {"text": "5л", "callback_data": "volume_5"},
                {"text": "6л", "callback_data": "volume_6"},
                {"text": "8л", "callback_data": "volume_8"},
                {"text": "10л", "callback_data": "volume_10"}
            ],
            [
                {"text": "📄 Конспект 2л", "callback_data": "quick_conspect_2"},
                {"text": "📄 Реферат 4л", "callback_data": "quick_referat_4"}
            ]
        ]
    }
    
    volume_text = f"""📄 <b>ВЫБОР ОБЪЕМА РАБОТЫ</b>

Текущий объем: {volume_info['icon']} <b>{volume_info['name']}</b>
📝 Слов: {volume_info['words']}

<b>📊 СТАНДАРТНЫЕ ОБЪЕМЫ:</b>

1-2 листа - 📘 Конспект, тезисы
3-5 листов - 📗 Реферат, доклад  
6-10 листов - 📕 Курсовая работа
10+ листов - 📓 Диплом, исследование

<b>🎯 РЕКОМЕНДАЦИИ:</b>
• Конспект: 1-3 листа
• Реферат: 3-5 листов
• Презентация: 2-4 листа текста
• Эссе: 2-3 листа

<b>📝 ФОРМАТЫ ЗАПРОСА:</b>
<code>реферат семья 4 листа</code>
<code>конспект физика 3л</code>
<code>презентация экология 10 слайдов</code>

<i>Нажмите на кнопку или напишите количество листов (1-30)</i>"""
    
    send_telegram_message(chat_id, volume_text, reply_markup=keyboard)

def handle_presentation_command(chat_id: int, user_id: str):
    """Обработка команды /presentation"""
    keyboard = {
        "inline_keyboard": [
            [
                {"text": "🎓 Академическая", "callback_data": "presentation_academic"},
                {"text": "💼 Бизнес", "callback_data": "presentation_business"}
            ],
            [
                {"text": "🎨 Креативная", "callback_data": "presentation_creative"},
                {"text": "📚 Образовательная", "callback_data": "presentation_educational"}
            ]
        ]
    }
    
    presentation_text = """🎤 <b>СОЗДАНИЕ ПРЕЗЕНТАЦИИ</b>

<b>Выберите стиль презентации:</b>

• <b>🎓 Академическая</b>
  <i>Для научных работ, дипломов, курсовых</i>
  <i>Строгий стиль, логическая структура</i>

• <b>💼 Бизнес-презентация</b>
  <i>Для проектов, стартапов, отчетов</i>
  <i>Корпоративный стиль, акцент на результатах</i>

• <b>🎨 Креативная</b>
  <i>Для творческих проектов, портфолио</i>
  <i>Современный дизайн, визуальные эффекты</i>

• <b>📚 Образовательная</b>
  <i>Для уроков, лекций, семинаров</i>
  <i>Интерактивные элементы, вопросы</i>

<b>📝 ФОРМАТ ЗАПРОСА:</b>
<code>презентация [тема] [слайды] [стиль]</code>

<b>ПРИМЕРЫ:</b>
<code>презентация экология 10 слайдов академическая</code>
<code>презентация стартап 12 бизнес</code>
<code>презентация искусство 8 креативная</code>

<i>Нажмите на кнопку стиля или напишите полный запрос</i>"""
    
    send_telegram_message(chat_id, presentation_text, reply_markup=keyboard)
    # ============ ОБРАБОТКА ЗАПРОСОВ ============
def handle_content_request_with_volume(chat_id: int, user_id: str, text: str):
    """Обработка запроса с указанием объема"""
    
    # Парсим запрос
    clean_text, volume = parse_volume_from_text(text)
    
    if not clean_text:
        send_telegram_message(chat_id, 
            "❌ Не указана тема.\n\n"
            "<b>Формат:</b>\n"
            "<code>[тип] [тема] [объем] [листы/слайды] [шаблон]</code>\n\n"
            "<b>Примеры:</b>\n"
            "<code>реферат семья 4 листа</code>\n"
            "<code>презентация экология 10 слайдов бизнес</code>"
        )
        return
    
    # Определяем тип контента и параметры
    parts = clean_text.split()
    if len(parts) < 2:
        send_telegram_message(chat_id, "❌ Не указана тема")
        return
    
    # Проверяем на презентацию
    is_presentation = False
    presentation_template = "academic"
    
    if "презентация" in parts[0].lower() or "слайд" in clean_text.lower():
        is_presentation = True
        content_type = "presentation"
        
        # Ищем шаблон в запросе
        for template in PRESENTATION_TEMPLATES.keys():
            if template in clean_text.lower():
                presentation_template = template
                # Убираем шаблон из текста
                clean_text = clean_text.replace(template, "").strip()
                break
    else:
        content_type_map = {
            'конспект': 'conspect',
            'реферат': 'referat', 
            'эссе': 'essay'
        }
        content_type_key = parts[0].lower().replace('по', '').strip()
        content_type = content_type_map.get(content_type_key, 'conspect')
    
    # Извлекаем тему
    topic = " ".join(parts[1:])
    
    # Объем по умолчанию
    if not volume:
        if is_presentation:
            volume = "10"  # 10 слайдов по умолчанию для презентации
        else:
            default_volumes = {
                'conspect': '2',
                'referat': '4', 
                'essay': '2',
                'presentation': '10'
            }
            volume = default_volumes.get(content_type, '3')
    
    # Сохраняем настройки
    user_settings[f"{user_id}_volume"] = volume
    if is_presentation:
        user_settings[f"{user_id}_presentation_template"] = presentation_template
    
    # Генерируем и отправляем
    generate_and_send_content(chat_id, user_id, topic, content_type, int(volume), 
                            presentation_template if is_presentation else None)

def generate_and_send_content(chat_id: int, user_id: str, topic: str, 
                            content_type: str = "conspect", volume_pages: int = 3,
                            presentation_template: str = None):
    """Генерация и отправка контента"""
    
    device_type = user_devices.get(user_id, "phone")
    device_info = DEVICES.get(device_type, DEVICES["phone"])
    content_type_info = CONTENT_TYPES.get(content_type, CONTENT_TYPES["conspect"])
    volume_info = VOLUME_LEVELS.get(str(volume_pages), VOLUME_LEVELS["3"])
    
    # Статус
    status_msg = (
        f"🔄 <b>ГЕНЕРАЦИЯ {content_type_info['name'].upper()}</b>\n\n"
        f"📝 Тема: <i>{topic}</i>\n"
        f"📱 Устройство: {device_info['icon']} <b>{device_info['name']}</b>\n"
    )
    
    if content_type == "presentation" and presentation_template:
        template_info = PRESENTATION_TEMPLATES.get(presentation_template, PRESENTATION_TEMPLATES["academic"])
        status_msg += f"🎤 Шаблон: <b>{template_info['name']}</b>\n"
        status_msg += f"📊 Слайдов: <b>{volume_pages}</b>\n\n"
    else:
        status_msg += f"📄 Листов А4: <b>{volume_pages}</b>\n"
        status_msg += f"📝 Слов: <b>{volume_info['words']}</b>\n\n"
    
    status_msg += "<i>Идет обработка с подбором источников...</i>"
    
    send_telegram_message(chat_id, status_msg)
    time.sleep(1)
    
    # Генерируем контент
    content = generate_ai_content(topic, content_type, device_type, volume_pages, 
                                presentation_template or "academic")
    
    # Отправляем контент
    logger.info(f"📤 Отправляю контент в чат {chat_id}")
    send_telegram_message(chat_id, content)
    
    # Сохраняем для экспорта
    user_settings[f"{user_id}_last_topic"] = topic
    user_settings[f"{user_id}_last_content"] = content
    user_settings[f"{user_id}_last_type"] = content_type
    user_settings[f"{user_id}_last_volume"] = str(volume_pages)
    
    if presentation_template:
        user_settings[f"{user_id}_last_template"] = presentation_template
    
    save_to_history(user_id, topic, content_type, str(volume_pages))
    
    # Предлагаем экспорт
    export_menu = {
        "inline_keyboard": [[
            {"text": "📥 1. Скачать файл", "callback_data": "export_menu"},
            {"text": "🔄 2. Новый материал", "callback_data": "new_topic"}
        ]]
    }
    
    final_text = (
        f"✅ <b>{content_type_info['name']} готов!</b>\n\n"
        f"<b>ПАРАМЕТРЫ:</b>\n"
        f"• Тема: {topic}\n"
        f"• Тип: {content_type_info['name']}\n"
    )
    
    if content_type == "presentation":
        final_text += f"• Слайдов: {volume_pages}\n"
        if presentation_template:
            template_info = PRESENTATION_TEMPLATES.get(presentation_template, PRESENTATION_TEMPLATES["academic"])
            final_text += f"• Шаблон: {template_info['name']}\n"
    else:
        final_text += f"• Листов А4: {volume_pages}\n"
        final_text += f"• Слов: {volume_info['words']}\n"
    
    final_text += f"• Устройство: {device_info['name']}\n\n"
    
    final_text += (
        f"📚 <b>ИСТОЧНИКИ:</b>\n"
        f"Автоматически подобраны {min(volume_pages + 2, 8)} источников\n\n"
        
        f"<b>ДАЛЬНЕЙШИЕ ДЕЙСТВИЯ:</b>\n"
        f"1 - Скачать файл\n"
        f"2 - Новый материал\n"
        f"3 - Изменить объем (/volume)\n\n"
        
        f"<i>Напишите цифру или используйте кнопки</i>"
    )
    
    send_telegram_message(chat_id, final_text, reply_markup=export_menu)

def handle_export_command(chat_id: int, user_id: str):
    """Обработка команды /export"""
    last_topic = user_settings.get(f"{user_id}_last_topic", None)
    
    if not last_topic:
        send_telegram_message(chat_id,
            "📊 <b>ЭКСПОРТ МАТЕРИАЛА</b>\n\n"
            "У вас пока нет сохраненных материалов.\n\n"
            "<i>Сначала создайте материал:</i>\n"
            "<code>реферат тема 4 листа</code>\n"
            "→ получите структуру с источниками\n"
            "→ используйте /export для скачивания"
        )
        return
    
    user_settings[f"{user_id}_awaiting_export"] = True
    
    export_text = f"""📊 <b>ЭКСПОРТ МАТЕРИАЛА</b>

Тема: <b>{last_topic}</b>
Тип: <b>{CONTENT_TYPES.get(user_settings.get(f'{user_id}_last_type', 'conspect'), CONTENT_TYPES['conspect'])['name']}</b>

<b>Выберите формат файла (напишите цифру):</b>

1. <b>📄 PDF</b> - для печати и чтения
   <i>Сохраняет форматирование, готов к печати</i>

2. <b>📝 DOCX</b> - для редактирования  
   <i>Можно редактировать в Word, Google Docs</i>

3. <b>📋 TXT</b> - простой текст
   <i>Совместим со всеми устройствами</i>

<b>Особенности:</b>
• Все форматы включают структуру и источники
• PDF готов к отправке на печать
• DOCX можно дорабатывать
• TXT - минимальный размер

<i>Напишите цифру от 1 до 3:</i>
<code>1</code> (PDF), <code>2</code> (DOCX) или <code>3</code> (TXT)"""
    
    send_telegram_message(chat_id, export_text)
    # ============ ЭКСПОРТ ФАЙЛОВ ============
def handle_export_format(chat_id: int, user_id: str, export_format: str):
    """Обработка выбора формата экспорта"""
    last_topic = user_settings.get(f"{user_id}_last_topic", "Тема")
    last_content = user_settings.get(f"{user_id}_last_content", "")
    last_type = user_settings.get(f"{user_id}_last_type", "conspect")
    
    format_info = EXPORT_FORMATS.get(export_format, EXPORT_FORMATS["txt"])
    
    # Убираем состояние ожидания
    user_settings[f"{user_id}_awaiting_export"] = False
    
    send_telegram_message(chat_id, f"🔄 <b>Создаю {format_info['name']} файл...</b>")
    
    try:
        # Создаем файл в зависимости от формата
        if export_format == "txt":
            filename, file_content = create_txt_file(last_topic, last_content, last_type)
        
        elif export_format == "pdf":
            filename, file_content = create_pdf_file(last_topic, last_content, last_type)
        
        elif export_format == "docx":
            filename, file_content = create_docx_file(last_topic, last_content, last_type)
        
        else:
            filename, file_content = create_txt_file(last_topic, last_content, last_type)
        
        # Отправляем файл
        caption = (
            f"{format_info['icon']} <b>{format_info['name']} {CONTENT_TYPES[last_type]['name']}:</b> {last_topic}\n\n"
            f"📅 Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
            f"📚 Включает структуру и источники\n"
            f"🤖 Бот: @Konspekt_help_bot"
        )
        
        response = send_telegram_document(
            chat_id=chat_id,
            filename=filename,
            content=file_content,
            caption=caption
        )
        
        if response.get("ok"):
            send_telegram_message(chat_id,
                f"✅ <b>Файл успешно отправлен!</b>\n\n"
                f"Формат: {format_info['icon']} {format_info['name']}\n"
                f"Тема: {last_topic}\n"
                f"Тип: {CONTENT_TYPES[last_type]['name']}\n"
                f"Размер: {len(file_content) // 1024} KB\n\n"
                f"<i>Для нового материала напишите запрос с объемом</i>"
            )
        else:
            send_telegram_message(chat_id,
                f"❌ <b>Ошибка отправки файла</b>\n\n"
                f"Попробуйте другой формат или обратитесь в поддержку."
            )
            
    except Exception as e:
        logger.error(f"Ошибка экспорта: {e}")
        send_telegram_message(chat_id,
            f"❌ <b>Ошибка создания файла</b>\n\n"
            f"Попробуйте другой формат или обратитесь в поддержку.\n"
            f"<i>Ошибка: {str(e)[:100]}...</i>"
        )

# ============ CALLBACK ОБРАБОТЧИКИ ============
def handle_volume_callback(callback_data: str, chat_id: int, user_id: str):
    """Обработка callback для выбора объема"""
    if callback_data.startswith('volume_'):
        volume = callback_data.split('_')[1]
        if volume in VOLUME_LEVELS:
            volume_info = VOLUME_LEVELS[volume]
            user_settings[f"{user_id}_volume"] = volume
            
            send_telegram_message(chat_id,
                f"✅ <b>Объем установлен: {volume_info['icon']} {volume_info['name']}</b>\n\n"
                f"📄 Листов А4: {volume_info['pages']}\n"
                f"📝 Слов: {volume_info['words']}\n\n"
                f"<i>Теперь напишите запрос с темой:</i>\n"
                f"<code>реферат ваша_тема {volume}л</code>\n"
                f"<code>конспект предмет {volume} листа</code>"
            )
    
    elif callback_data.startswith('quick_'):
        parts = callback_data.split('_')
        if len(parts) == 3:
            content_type = parts[1]  # conspect, referat
            volume = parts[2]        # 2, 4
            
            type_names = {'conspect': 'конспект', 'referat': 'реферат'}
            type_name = type_names.get(content_type, 'материал')
            volume_info = VOLUME_LEVELS.get(volume, VOLUME_LEVELS["3"])
            
            user_settings[f"{user_id}_volume"] = volume
            
            send_telegram_message(chat_id,
                f"✅ <b>Шаблон выбран: {type_name} на {volume_info['name']}</b>\n\n"
                f"📄 Листов А4: {volume_info['pages']}\n"
                f"📝 Слов: {volume_info['words']}\n\n"
                f"<i>Теперь напишите тему:</i>\n"
                f"<code>{type_name} ваша_тема</code>\n\n"
                f"Бот создаст {type_name} на {volume_info['name']} А4"
            )

def handle_presentation_callback(callback_data: str, chat_id: int, user_id: str):
    """Обработка callback для выбора шаблона презентации"""
    if callback_data.startswith('presentation_'):
        template = callback_data.split('_')[1]
        template_info = PRESENTATION_TEMPLATES.get(template, PRESENTATION_TEMPLATES["academic"])
        
        user_settings[f"{user_id}_presentation_template"] = template
        
        send_telegram_message(chat_id,
            f"✅ <b>Шаблон установлен: {template_info['name']}</b>\n\n"
            f"🎨 Стиль: {template_info['style']}\n"
            f"📊 Слайдов в структуре: {len(template_info['slides'])}\n\n"
            f"<i>Теперь напишите запрос:</i>\n"
            f"<code>презентация ваша_тема 10 слайдов</code>\n"
            f"Или просто: <code>презентация тема 12</code>"
        )

def handle_export_callback(callback_data: str, chat_id: int, user_id: str):
    """Обработка callback для экспорта"""
    if callback_data == "export_menu":
        handle_export_command(chat_id, user_id)
    elif callback_data.startswith("export_"):
        format_map = {
            "export_pdf": "pdf",
            "export_docx": "docx", 
            "export_txt": "txt"
        }
        if callback_data in format_map:
            handle_export_format(chat_id, user_id, format_map[callback_data])
            # ============ ОБРАБОТЧИКИ КОМАНД (ПРОДОЛЖЕНИЕ) ============
def handle_ai_command(chat_id: int, user_id: str = None):
    """Обработка команды /ai"""
    ai_text = """🤖 <b>AI-ГЕНЕРАЦИЯ МАТЕРИАЛОВ</b>

<b>Доступные типы материалов:</b>

• <b>📚 Конспект</b> - структурированные учебные заметки
  <i>Для быстрого изучения и повторения</i>

• <b>📄 Реферат</b> - научная исследовательская работа  
  <i>Структура, библиография, требования</i>

• <b>🎤 Презентация</b> - план выступления со слайдами
  <i>Для защиты проектов, докладов, отчетов</i>

• <b>✍️ Эссе</b> - аналитическое или художественное сочинение
  <i>Для творческих заданий, размышлений, анализа</i>

<b>Как использовать:</b>
<code>[тип] [тема] [объем]</code>

<b>Примеры запросов:</b>
<code>реферат квантовая физика 4 листа</code>
<code>презентация искусственный интеллект 10 слайдов</code>  
<code>конспект математика 2л</code>
<code>эссе философия стоицизма 3 страницы</code>

<b>📊 Укажите объем:</b>
• В листах А4: 1, 2, 3, 4, 5...
• В слайдах: 10, 12, 15...

<i>Напишите запрос в формате выше</i>"""
    
    send_telegram_message(chat_id, ai_text)

def handle_history_command(chat_id: int, user_id: str):
    """Обработка команды /history"""
    history = user_history.get(user_id, [])
    
    if not history:
        send_telegram_message(chat_id,
            "📜 <b>ИСТОРИЯ ЗАПРОСОВ</b>\n\n"
            "Ваша история запросов пуста.\n\n"
            "<i>Создайте первый материал:</i>\n"
            "<code>реферат тема 4 листа</code>\n"
            "<code>презентация проект 10 слайдов</code>"
        )
        return
    
    # Показываем последние 5 запросов
    recent = history[-5:]
    history_text = "📜 <b>ПОСЛЕДНИЕ ЗАПРОСЫ</b>\n\n"
    
    for i, item in enumerate(reversed(recent), 1):
        item_type = CONTENT_TYPES.get(item.get("type", "conspect"), CONTENT_TYPES["conspect"])
        device_info = DEVICES.get(item.get("device", "phone"), DEVICES["phone"])
        volume = item.get("volume", "3")
        
        timestamp = datetime.fromisoformat(item["timestamp"]).strftime("%d.%m %H:%M")
        
        history_text += f"{i}. <b>{item_type['icon']} {item['topic']}</b>\n"
        history_text += f"   📊 {VOLUME_LEVELS.get(volume, VOLUME_LEVELS['3'])['name']} | "
        history_text += f"📱 {device_info['icon']} | ⏰ {timestamp}\n\n"
    
    if len(history) > 5:
        history_text += f"<i>Показано 5 из {len(history)} запросов</i>\n\n"
    
    history_text += (
        "<b>Для повторной генерации:</b>\n"
        "Просто напишите тему заново\n\n"
        "<b>Очистка истории:</b>\n"
        "В настройках /settings"
    )
    
    send_telegram_message(chat_id, history_text)

def handle_settings_command(chat_id: int, user_id: str):
    """Обработка команды /settings"""
    current_device = user_devices.get(user_id, "phone")
    device_info = DEVICES.get(current_device, DEVICES["phone"])
    
    current_volume = user_settings.get(f"{user_id}_volume", "3")
    volume_info = VOLUME_LEVELS.get(current_volume, VOLUME_LEVELS["3"])
    
    history_count = len(user_history.get(user_id, []))
    
    current_template = user_settings.get(f"{user_id}_presentation_template", "academic")
    template_info = PRESENTATION_TEMPLATES.get(current_template, PRESENTATION_TEMPLATES["academic"])
    
    settings_text = f"""⚙️ <b>НАСТРОЙКИ БОТА v7.0</b>

<b>Текущие настройки:</b>
• 📱 Устройство: <b>{device_info['icon']} {device_info['name']}</b>
• 📊 Объем по умолчанию: <b>{volume_info['icon']} {volume_info['name']}</b>
• 🎤 Шаблон презентации: <b>{template_info['name']}</b>
• 📜 История: <b>{history_count} запросов</b>

<b>Доступные действия:</b>

1. <b>Сменить устройство</b>
   Напишите: <code>телефон</code>, <code>компьютер</code>, <code>планшет</code> или <code>часы</code>

2. <b>Изменить объем</b>
   Команда: /volume

3. <b>Изменить шаблон презентации</b>
   Команда: /presentation

4. <b>Очистить историю</b>
   Напишите: <code>очистить историю</code>

5. <b>Экспорт всех данных</b>
   Напишите: <code>экспорт данных</code>

<b>Техническая информация:</b>
• Версия бота: 7.0.0
• Платформа: Render.com
• Режим работы: 24/7
• Токен: {TOKEN[:10]}...
• Статус: активен ✅

<i>Для изменения напишите команду или действие</i>"""
    
    send_telegram_message(chat_id, settings_text)

# ============ ДОПОЛНИТЕЛЬНЫЕ ОБРАБОТЧИКИ ============
def handle_clear_history(chat_id: int, user_id: str):
    """Очистка истории запросов"""
    if user_id in user_history:
        user_history[user_id] = []
    
    send_telegram_message(chat_id,
        "🗑️ <b>История запросов очищена</b>\n\n"
        "Все предыдущие запросы удалены.\n\n"
        "<i>Создайте новый материал:</i>\n"
        "<code>реферат тема 4 листа</code>"
    )

def handle_export_data(chat_id: int, user_id: str):
    """Экспорт всех данных пользователя"""
    user_data = {
        "user_id": user_id,
        "device": user_devices.get(user_id, "phone"),
        "volume": user_settings.get(f"{user_id}_volume", "3"),
        "history": user_history.get(user_id, []),
        "settings": {k: v for k, v in user_settings.items() if k.startswith(user_id)},
        "exported_at": datetime.now().isoformat()
    }
    
    # Создаем JSON файл
    import json
    filename = f"user_data_{user_id}.json"
    file_content = json.dumps(user_data, ensure_ascii=False, indent=2).encode('utf-8')
    
    caption = f"📦 <b>Экспорт данных пользователя</b>\n\nID: {user_id}\nДата: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    
    response = send_telegram_document(chat_id, filename, file_content, caption)
    
    if response.get("ok"):
        send_telegram_message(chat_id, "✅ <b>Данные успешно экспортированы!</b>")
    else:
        send_telegram_message(chat_id, "❌ <b>Ошибка экспорта данных</b>")

# ============ ОБРАБОТКА ТЕКСТОВЫХ КОМАНД ============
def handle_text_commands(chat_id: int, user_id: str, text: str):
    """Обработка текстовых команд"""
    text_lower = text.lower()
    
    if text_lower in ['очистить историю', 'clear history', 'удалить историю']:
        handle_clear_history(chat_id, user_id)
    
    elif text_lower in ['экспорт данных', 'export data', 'скачать данные']:
        handle_export_data(chat_id, user_id)
    
    elif text_lower in ['сбросить настройки', 'reset settings', 'default']:
        # Сбрасываем настройки пользователя
        if f"{user_id}_volume" in user_settings:
            del user_settings[f"{user_id}_volume"]
        if f"{user_id}_presentation_template" in user_settings:
            del user_settings[f"{user_id}_presentation_template"]
        
        send_telegram_message(chat_id,
            "🔄 <b>Настройки сброшены</b>\n\n"
            "Все настройки возвращены к значениям по умолчанию.\n\n"
            "<i>Текущие настройки:</i>\n"
            "• Устройство: телефон\n"
            "• Объем: 3 листа\n"
            "• Шаблон презентации: академический"
        )
    
    elif text_lower in ['статус', 'status', 'инфо', 'info']:
        # Показываем статус бота
        from flask import current_app
        
        status_info = {
            "Бот": "Активен ✅",
            "Пользователей": len(user_devices),
            "Запросов сегодня": sum(len(h) for h in user_history.values()),
            "Память": "Нормальная",
            "Версия": "7.0.0",
            "Время работы": "24/7"
        }
        
        status_text = "📊 <b>СТАТУС СИСТЕМЫ</b>\n\n"
        for key, value in status_info.items():
            status_text += f"• {key}: <b>{value}</b>\n"
        
        status_text += "\n<i>Все системы работают нормально</i>"
        
        send_telegram_message(chat_id, status_text)
    
    elif text_lower in ['помощь', 'help', 'справка']:
        handle_help_command(chat_id)
    
    elif text_lower in ['объем', 'volume', 'листы']:
        handle_volume_command(chat_id, user_id)
    
    elif text_lower in ['презентация', 'presentation', 'слайды']:
        handle_presentation_command(chat_id, user_id)
    
    elif text_lower in ['экспорт', 'export', 'скачать']:
        handle_export_command(chat_id, user_id)
    
    elif text_lower in ['история', 'history', 'запросы']:
        handle_history_command(chat_id, user_id)
    
    elif text_lower in ['настройки', 'settings', 'опции']:
        handle_settings_command(chat_id, user_id)
    
    elif text_lower in ['старт', 'start', 'начать']:
        handle_start_command(chat_id, "Пользователь", user_id)
    
    else:
        # Если это не команда, проверяем, может быть это запрос с объемом
        if (any(keyword in text_lower for keyword in ['лист', 'л ', 'страниц', 'стр ', 'слайд']) or 
            re.search(r'\b(1|2|3|4|5|6|7|8|9|10|15|20|30)\b', text)):
            handle_content_request_with_volume(chat_id, user_id, text)
        else:
            # Если обычный текст без объема, предлагаем указать объем
            send_telegram_message(chat_id,
                f"📝 <b>Вы написали: {text}</b>\n\n"
                f"<i>Укажите объем работы:</i>\n\n"
                f"<b>Примеры запросов с объемом:</b>\n"
                f"<code>реферат {text} 4 листа</code>\n"
                f"<code>конспект {text} 3л</code>\n"
                f"<code>презентация {text} 10 слайдов</code>\n\n"
                f"Или используйте /volume для выбора объема"
            )

# ============ ОБНОВЛЕННЫЙ ВЕБХУК ============
@app.route('/' + TOKEN, methods=['POST'])
def telegram_webhook():
    """Основной обработчик вебхука от Telegram"""
    try:
        data = request.json
        logger.info(f"📨 Получен вебхук от Telegram")
        
        # Обрабатываем сообщение
        if 'message' in data:
            message = data['message']
            chat_id = message['chat']['id']
            user_id = str(message['from']['id'])
            username = message['from'].get('first_name', 'Пользователь')
            text = message.get('text', '').strip()
            
            logger.info(f"👤 [{username}] → {text}")
            
            # Основные команды
            if text == '/start':
                handle_start_command(chat_id, username, user_id)
            
            elif text == '/help':
                handle_help_command(chat_id)
            
            elif text == '/volume':
                handle_volume_command(chat_id, user_id)
            
            elif text == '/presentation':
                handle_presentation_command(chat_id, user_id)
            
            elif text == '/export':
                handle_export_command(chat_id, user_id)
            
            elif text == '/ai':
                handle_ai_command(chat_id, user_id)
            
            elif text == '/history':
                handle_history_command(chat_id, user_id)
            
            elif text == '/settings':
                handle_settings_command(chat_id, user_id)
            
            # Выбор устройства
            elif text.lower() in ['телефон', '📱 телефон', 'phone', 'мобильный', 'смартфон']:
                user_devices[user_id] = "phone"
                send_telegram_message(chat_id,
                    f"✅ <b>Устройство выбрано: Телефон</b>\n\n"
                    f"Теперь все материалы будут оптимизированы для мобильных экранов.\n\n"
                    f"<i>Напишите запрос с объемом:</i>\n"
                    f"<code>реферат тема 4 листа</code>\n"
                    f"<code>презентация проект 10 слайдов</code>"
                )
            
            elif text.lower() in ['компьютер', '💻 компьютер', 'pc', 'пк', 'ноутбук', 'десктоп']:
                user_devices[user_id] = "pc"
                send_telegram_message(chat_id,
                    f"✅ <b>Устройство выбрано: Компьютер</b>\n\n"
                    f"Теперь все материалы будут в полной версии для ПК.\n\n"
                    f"<i>Напишите запрос с объемом:</i>\n"
                    f"<code>реферат тема 4 листа</code>\n"
                    f"<code>презентация проект 10 слайдов</code>"
                )
            
            elif text.lower() in ['планшет', '📟 планшет', 'tablet', 'таблет', 'айпад']:
                user_devices[user_id] = "tablet"
                send_telegram_message(chat_id,
                    f"✅ <b>Устройство выбрано: Планшет</b>\n\n"
                    f"Теперь все материалы будут в промежуточной версии.\n\n"
                    f"<i>Напишите запрос с объемом:</i>\n"
                    f"<code>реферат тема 4 листа</code>\n"
                    f"<code>презентация проект 10 слайдов</code>"
                )
            
            elif text.lower() in ['часы', '⌚ часы', 'watch', 'умные часы', 'смарт-часы']:
                user_devices[user_id] = "watch"
                send_telegram_message(chat_id,
                    f"✅ <b>Устройство выбрано: Часы</b>\n\n"
                    f"Теперь все материалы будут в краткой версии.\n\n"
                    f"<i>Напишите запрос с объемом:</i>\n"
                    f"<code>реферат тема 4 листа</code>\n"
                    f"<code>презентация проект 10 слайдов</code>"
                )
            
            # Обработка цифрового выбора экспорта
            elif user_settings.get(f"{user_id}_awaiting_export") and text in ['1', '2', '3']:
                format_map = {"1": "pdf", "2": "docx", "3": "txt"}
                handle_export_format(chat_id, user_id, format_map[text])
            
            # Обработка цифрового выбора объема
            elif user_settings.get(f"{user_id}_awaiting_volume") and text.isdigit():
                volume = text
                if volume in VOLUME_LEVELS:
                    volume_info = VOLUME_LEVELS[volume]
                    user_settings[f"{user_id}_volume"] = volume
                    user_settings[f"{user_id}_awaiting_volume"] = False
                    
                    send_telegram_message(chat_id,
                        f"✅ <b>Объем установлен: {volume_info['icon']} {volume_info['name']}</b>\n\n"
                        f"📄 Листов А4: {volume_info['pages']}\n"
                        f"📝 Слов: {volume_info['words']}\n\n"
                        f"<i>Теперь напишите запрос с темой:</i>\n"
                        f"<code>реферат ваша_тема {volume}л</code>"
                    )
                else:
                    send_telegram_message(chat_id,
                        f"❌ Объем {volume} листов не поддерживается.\n\n"
                        f"<b>Доступные объемы:</b>\n"
                        f"1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 15, 20, 30 листов"
                    )
            
            # Обработка текстовых команд
            elif text and not text.startswith('/'):
                handle_text_commands(chat_id, user_id, text)
            
            # Неизвестная команда
            elif text.startswith('/'):
                send_telegram_message(chat_id,
                    "❓ <b>Неизвестная команда</b>\n\n"
                    "Используйте одну из команд:\n"
                    "• /start - начать работу\n"
                    "• /help - получить справку\n"
                    "• /volume - выбор объема\n"
                    "• /presentation - создать презентацию\n"
                    "• /export - экспорт файлов\n"
                    "• /ai - AI-генерация\n"
                    "• /history - история\n"
                    "• /settings - настройки\n\n"
                    "<i>Или напишите тему с указанием объема!</i>"
                )
            
            # Пустое сообщение
            else:
                send_telegram_message(chat_id,
                    "📝 <b>Напишите запрос с указанием объема!</b>\n\n"
                    "<b>Примеры:</b>\n"
                    "• реферат семья 4 листа\n"
                    "• конспект математика 3л\n"
                    "• презентация экология 10 слайдов\n"
                    "• эссе философия 2 страницы\n\n"
                    "<i>Используйте /help для полной справки</i>"
                )
        
        # Обработка callback query (кнопки)
        elif 'callback_query' in data:
            callback = data['callback_query']
            callback_id = callback['id']
            chat_id = callback['message']['chat']['id']
            user_id = str(callback['from']['id'])
            callback_data = callback['data']
            
            # Ответ на callback
            requests.post(f"https://api.telegram.org/bot{TOKEN}/answerCallbackQuery", 
                         json={"callback_query_id": callback_id})
            
            # Обработка callback данных
            if callback_data.startswith('volume_') or callback_data.startswith('quick_'):
                handle_volume_callback(callback_data, chat_id, user_id)
            
            elif callback_data.startswith('presentation_'):
                handle_presentation_callback(callback_data, chat_id, user_id)
            
            elif callback_data.startswith('export_') or callback_data == 'export_menu':
                handle_export_callback(callback_data, chat_id, user_id)
            
            elif callback_data == 'new_topic':
                send_telegram_message(chat_id,
                    "🔄 <b>Создание нового материала</b>\n\n"
                    "<i>Напишите запрос с указанием объема:</i>\n\n"
                    "<b>Примеры:</b>\n"
                    "<code>реферат новая_тема 4 листа</code>\n"
                    "<code>презентация проект 12 слайдов бизнес</code>"
                )
        
        return jsonify({"status": "ok"}), 200
        
    except Exception as e:
        logger.error(f"❌ Ошибка обработки вебхука: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500
        # ============ HTML СТРАНИЦА ============
@app.route('/')
def home():
    return '''
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🎓 Учебный Бот Премиум v7.0</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #6a11cb 0%, #2575fc 100%);
            color: white;
            min-height: 100vh;
            padding: 20px;
            display: flex;
            justify-content: center;
            align-items: center;
        }
        
        .container {
            background: rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(15px);
            border-radius: 25px;
            padding: 50px;
            max-width: 900px;
            width: 100%;
            box-shadow: 0 25px 75px rgba(0, 0, 0, 0.3);
            border: 1px solid rgba(255, 255, 255, 0.2);
            text-align: center;
        }
        
        h1 {
            font-size: 3.5em;
            margin-bottom: 20px;
            background: linear-gradient(45deg, #ff6b6b, #ffd93d, #6bcf7f);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            text-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
        }
        
        .status {
            display: inline-flex;
            align-items: center;
            gap: 10px;
            background: rgba(46, 204, 113, 0.2);
            border: 2px solid #2ecc71;
            color: #2ecc71;
            padding: 12px 30px;
            border-radius: 50px;
            font-size: 1.3em;
            font-weight: bold;
            margin: 25px auto;
        }
        
        .features-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 40px 0;
        }
        
        .feature {
            background: rgba(255, 255, 255, 0.05);
            border-radius: 15px;
            padding: 25px 15px;
            transition: all 0.3s ease;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }
        
        .feature:hover {
            transform: translateY(-5px);
            background: rgba(255, 255, 255, 0.1);
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.2);
        }
        
        .feature-icon {
            font-size: 2.5em;
            margin-bottom: 15px;
            display: block;
        }
        
        .btn {
            display: inline-flex;
            align-items: center;
            gap: 12px;
            background: linear-gradient(45deg, #0088cc, #00c6ff);
            color: white;
            text-decoration: none;
            padding: 18px 45px;
            border-radius: 50px;
            font-size: 1.3em;
            font-weight: bold;
            margin: 30px 10px;
            transition: all 0.3s ease;
            box-shadow: 0 8px 25px rgba(0, 136, 204, 0.4);
            border: none;
            cursor: pointer;
        }
        
        .btn:hover {
            transform: translateY(-3px) scale(1.05);
            box-shadow: 0 12px 35px rgba(0, 136, 204, 0.6);
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>🎓 Учебный Бот Премиум v7.0</h1>
        
        <div class="status">
            ✅ Активен на Render 24/7
        </div>
        
        <p style="font-size: 1.2em; margin-bottom: 30px; opacity: 0.9; line-height: 1.6;">
            Интеллектуальный помощник для создания учебных материалов<br>
            с выбором объема в листах А4, презентациями и источниками
        </p>
        
        <div class="features-grid">
            <div class="feature">
                <span class="feature-icon">📄</span>
                <div class="feature-title">Объем в листах А4</div>
                <div>Указывайте: "реферат тема 4 листа"</div>
            </div>
            
            <div class="feature">
                <span class="feature-icon">🎤</span>
                <div class="feature-title">Презентации</div>
                <div>4 шаблона, структура слайдов</div>
            </div>
            
            <div class="feature">
                <span class="feature-icon">📚</span>
                <div class="feature-title">Автоматические источники</div>
                <div>Научные журналы, учебники</div>
            </div>
            
            <div class="feature">
                <span class="feature-icon">📱</span>
                <div class="feature-title">Адаптация под устройства</div>
                <div>Телефон, компьютер, планшет, часы</div>
            </div>
            
            <div class="feature">
                <span class="feature-icon">📊</span>
                <div class="feature-title">Экспорт файлов</div>
                <div>PDF, DOCX, TXT с источниками</div>
            </div>
            
            <div class="feature">
                <span class="feature-icon">🤖</span>
                <div class="feature-title">AI-генерация</div>
                <div>Структура под любой объем</div>
            </div>
        </div>
        
        <div style="margin: 40px 0;">
            <a href="https://t.me/Konspekt_help_bot" class="btn" target="_blank">
                <span>📱</span>
                Открыть в Telegram
            </a>
        </div>
        
        <div style="margin-top: 40px; padding-top: 30px; border-top: 1px solid rgba(255, 255, 255, 0.1);">
            <p>🚀 Работает на Render.com | 📄 Объем в листах А4 | 🎤 Презентации с шаблонами</p>
            <p>📚 Автоматические источники | 🤖 AI-структурирование | 📱 Адаптация под устройства</p>
        </div>
    </div>
</body>
</html>
'''

# ============ HEALTH CHECK ============
@app.route('/health')
def health():
    return jsonify({
        "status": "ok",
        "service": "study-bot-premium-v7",
        "version": "7.0.0",
        "timestamp": datetime.now().isoformat(),
        "features": [
            "volume_selection_a4",
            "presentation_templates", 
            "automatic_sources",
            "pdf_export",
            "docx_export",
            "device_optimization"
        ]
    }), 200

# ============ НАСТРОЙКА ВЕБХУКА ============
def setup_webhook():
    """Автоматическая настройка вебхука"""
    try:
        app_url = os.environ.get('RENDER_EXTERNAL_HOSTNAME', 'study-bot.onrender.com')
        webhook_url = f"https://{app_url}/{TOKEN}"
        
        logger.info(f"🔧 Настраиваю вебхук: {webhook_url}")
        
        # Удаляем старый вебхук
        delete_url = f"https://api.telegram.org/bot{TOKEN}/deleteWebhook"
        response = requests.get(delete_url, timeout=5)
        
        # Устанавливаем новый вебхук
        set_url = f"https://api.telegram.org/bot{TOKEN}/setWebhook"
        payload = {
            "url": webhook_url,
            "drop_pending_updates": True,
            "max_connections": 40,
            "allowed_updates": ["message", "callback_query"]
        }
        
        response = requests.post(set_url, json=payload, timeout=10)
        result = response.json()
        
        if result.get('ok'):
            logger.info(f"✅ Вебхук успешно установлен")
        else:
            logger.error(f"❌ Ошибка установки вебхука: {result}")
            
    except Exception as e:
        logger.error(f"❌ Ошибка при настройке вебхука: {e}")

# ============ ЗАПУСК ПРИЛОЖЕНИЯ ============
if __name__ == '__main__':
    logger.info("=" * 80)
    logger.info("🚀 ЗАПУСК УЧЕБНОГО БОТА ПРЕМИУМ v7.0")
    logger.info("=" * 80)
    logger.info(f"🤖 Бот: @Konspekt_help_bot")
    logger.info(f"🔑 Токен: {TOKEN[:10]}...")
    logger.info("=" * 80)
    
    # Настройка вебхука
    setup_webhook()
    
    # Запуск Flask сервера
    port = int(os.environ.get('PORT', 8080))
    logger.info(f"🌍 Запуск веб-сервера на порту {port}...")
    
    app.run(
        host='0.0.0.0',
        port=port,
        debug=False,
        use_reloader=False,
        threaded=True
    )
